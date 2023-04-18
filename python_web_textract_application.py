from collections import defaultdict
import streamlit as st
import boto3
import io
from PIL import Image
import io
import docx

st.title("AWS Textract Text Extraction App")

aws_management_console = boto3.session.Session(profile_name='USER_NAME')   #replace USER_NAME with your IAM user 
client = aws_management_console.client(service_name = 'textract', region_name = 'REGION')  #replace REGION with your IAM user's region



# Define a function to extract text from an image
def extract_text_from_image(image):
    response = client.detect_document_text(Document={"Bytes": image})

    # Extract the text from the response
    text = ""

    for item in response["Blocks"]:
        if item["BlockType"] == "LINE":
            text += item["Text"] + "\n"

    # Display the extracted text
    st.subheader("Extracted Text")
    st.write(text)
    download_text(text)

def download_text(text):
    doc = docx.Document()
    doc.add_paragraph(text)

    output_file = io.BytesIO()
    doc.save(output_file)
    output_file.seek(0)
    st.download_button(
        label="Download Doc File",
        data=output_file,
        file_name="extracted_text.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# Define a function to extract text from a PDF file
def extract_text_from_pdf(pdf_file):
    with io.BytesIO(pdf_file) as pdf_buffer:
        response = client.detect_document_text(Document={'Bytes': pdf_buffer.read()})
        text_blocks = [item['Text'] for item in response['Blocks'] if item['BlockType'] == "WORD"]
        if len(text_blocks) > 0:
            text = "\n".join(text_blocks)
            st.success("Text extracted from pdf:")
            st.write(text)
            download_text(text)
        else:
            st.warning("No text found in the image.")
        
        
        
# Define a function to extract text from a form
def extract_text_from_form(form):
    key_map, value_map, block_map = get_ks_map(form)
    
    kvs = get_kv_relationship(key_map, value_map, block_map)

    for key, value in kvs.items():
        st.write(key,":",value)

    download_text_from_form(kvs)

def download_text_from_form(kvs):
    doc = docx.Document()

    for key, value in kvs.items():
        doc.add_paragraph(f"{key}: {value}")

    output_file = io.BytesIO()
    doc.save(output_file)
    output_file.seek(0)
    st.download_button(
        label="Download Doc File",
        data=output_file,
        file_name="extracted_text.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def get_kv_relationship(key_map, value_map, block_map):
    kvs = defaultdict(list)
    for block_id, key_block in key_map.items():
        value_block = find_value_block(key_block, value_map)
        key = get_text(key_block, block_map)
        val = get_text(value_block, block_map)
        kvs[key].append(val)
    return kvs

def find_value_block(key_block, value_map):
    for relationship in key_block['Relationships']:
        if relationship['Type'] == 'VALUE':
            for value_id in relationship['Ids']:
                value_block = value_map[value_id]
    return value_block

def get_text(result, blocks_map):
    text=''
    if 'Relationships' in result:
        for relationship in result['Relationships']:
            if relationship['Type'] == 'CHILD':
                for child_id in relationship['Ids']:
                    word = blocks_map[child_id]
                    if word['BlockType'] == 'WORD':
                        text+=word['Text'] + ' '
                    if word['BlockType'] == 'SELECTION_ELEMENT':
                        if word['SelectionStatus'] == 'SELECTED':
                            text+='X'
    return text

def get_ks_map(img_to_bytes):
    response = client.analyze_document(Document={'Bytes': img_to_bytes, }, FeatureTypes=['FORMS'])
    blocks = response['Blocks']

    key_map = {}
    value_map = {}
    block_map = {}
    for block in blocks:
        block_id = block['Id']
        block_map[block_id] = block
        if block['BlockType'] == "KEY_VALUE_SET":
            if 'KEY' in block['EntityTypes']:
                key_map[block_id] = block
            else:
                value_map[block_id] = block

    return key_map, value_map, block_map

# Define the main function that will be called when the "Extract Text" button is clicked
def extract_text(file, file_type):
    if file_type == "Image":
        extracted_text = extract_text_from_image(file.read())
        img = Image.open(file)
        st.image(img, caption="Uploaded Image")
    elif file_type == "PDF":
        extracted_text = extract_text_from_pdf(file.read())
        st.write(extracted_text)
    elif file_type == "Form":
        extracted_text = extract_text_from_form(file.read())
        st.success("Text extracted from form:")
        img = Image.open(file)
        st.image(img, caption="Uploaded Image")

# Create an option menu for selecting the file type
file_type = st.selectbox("Select file type", options=["Image", "PDF", "Form"])

# Create a file uploader
file = st.file_uploader("Upload file", type=["jpg", "jpeg", "png", "pdf"])

# Create a button to extract text from the uploaded file
if st.button("Extract Text"):
    if file is not None:
        extract_text(file, file_type)
    else:
        st.write("Please upload a file.")
