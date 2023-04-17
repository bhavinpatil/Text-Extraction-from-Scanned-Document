from base64 import encode
from collections import defaultdict
from lib2to3.pytree import Node
from msilib.schema import Feature
from platform import node
from pydoc import Doc
import tkinter as tk
from tkinter import filedialog
from tkinter.filedialog import askopenfile, asksaveasfilename
from xml.dom.minidom import Document
import boto3
from PIL import Image, ImageTk
import docx

my_window = tk.Tk()
my_window.geometry("1000x800")
my_window.title("AWS Textract")
doc_type = tk.StringVar(my_window)
doc_type.set("JPG Image")

l1 = tk.Label(my_window, text="Upload an Image", width=30, font=('times', 18, 'bold'))
l1.pack()

w = tk.OptionMenu(my_window, doc_type, "JPG Image", "Form", "Table", "Signature", "Math Equations")
w.pack()

b1 = tk.Button(my_window, text="Upload File", width=15, command=lambda: upload_file())
b1.pack()


img_label = None
save_as = None

aws_management_console = boto3.session.Session(profile_name='USER_NAME')   #replace USER_NAME with your IAM user 
client = aws_management_console.client(service_name = 'textract', region_name = 'REGION')  #replace REGION with your IAM user's region


def upload_file():
    global img_label
    global save_as
    
    if save_as is not Node:
        if isinstance(save_as, tk.Button):
            save_as.destroy()


    global img

    if doc_type.get() == "JPG Image":
        f_types = [('JPG Files', '*jpg')]  # JPG Images
    elif doc_type.get() == "Form":
        f_types = [('Form Files', '*.jpg')]  # Form Files
    elif doc_type.get() == "Table":
        f_types = [('Jpg Files', '*.jpg')]  # Table from JPG Images
    elif doc_type.get() == "Signature":
        f_types = [('Jpg Files', '*.jpg')]  # Signatures from JPG Image
    elif doc_type.get() == "Math Equations":
        f_types = [('Jph Files', '*.jpg')]  # Math Equations from JPG Image

    filename = filedialog.askopenfilename(filetypes=f_types)  # dialog box for taking the image

    response = None
    img = Image.open(filename)
    # resizing the uploaded image with respect to window size
    img_resize = img.resize((750, 650))
    img = ImageTk.PhotoImage(img_resize)
    img_to_bytes = get_image_byte(filename)

    if img_label is not Node:
        if isinstance(img_label, tk.Label):
            img_label.destroy()
        
    img_label = tk.Label(my_window, image=img)
    img_label.pack()

# ---------------------------------------------------------------------------------
    if doc_type.get() == "JPG Image":

        response = client.detect_document_text(Document={'Bytes': img_to_bytes, })
        extracted_text = ""
        with open('output.txt', 'w') as f:
            for item in response['Blocks']:
                if item['BlockType'] == 'LINE':
                    print(item['Text'], file=f)
                    print(item['Text'])
                    extracted_text += item['Text'] + "\n"
        save_as = tk.Button(my_window, text="Save as", width=15, command=lambda:save_file(extracted_text))
        save_as.pack()


# ---------------------------------------------------------------------------------
    if doc_type.get() == "Form":

        key_map, value_map, block_map = get_ks_map(img_to_bytes)

        kvs = get_kv_relationship(key_map, value_map, block_map)

        with open('output.txt', 'w') as f:
            for key, value in kvs.items():
                print(key,":",value, file=f)
                print(key,":",value)
        
        save_as = tk.Button(my_window, text="Save as", width=15, command=lambda:save_to_doc(kvs))
        save_as.pack()

# ---------------------------------------------------------------------------------
    if doc_type.get() == "Table":
        
        
# ---------------------------------------------------------------------------------
    # if doc_type.get() == "Signature":

# ---------------------------------------------------------------------------------
    # if doc_type.get() == "Math Equations":
# ---------------------------------------------------------------------------------

def save_to_doc(kvs):
    document = docx.Document()
    for key, value in kvs.items():
        document.add_paragraph(f"{key}: {value}")
    filename = filedialog.asksaveasfilename(defaultextension='.docx')
    document.save(filename)

def get_kv_relationship(key_map, value_map, block_map):
    kvs = defaultdict(list)
    for block_id, key_block in key_map.items():
        value_block = find_value_block(key_block, value_map)
        key = get_text(key_block, block_map)
        val = get_text(value_block, block_map)
        kvs[key].append(val)
    return kvs

def find_value_block(key_block, value_map):
    for relationship in key_block['Relationships']:
        if relationship['Type'] == 'VALUE':
            for value_id in relationship['Ids']:
                value_block = value_map[value_id]
    return value_block

def get_text(result, blocks_map):
    text=''
    if 'Relationships' in result:
        for relationship in result['Relationships']:
            if relationship['Type'] == 'CHILD':
                for child_id in relationship['Ids']:
                    word = blocks_map[child_id]
                    if word['BlockType'] == 'WORD':
                        text+=word['Text'] + ' '
                    if word['BlockType'] == 'SELECTION_ELEMENT':
                        if word['SelectionStatus'] == 'SELECTED':
                            text+='X'
    return text

def get_ks_map(img_to_bytes):
    response = client.analyze_document(Document={'Bytes': img_to_bytes, }, FeatureTypes=['FORMS'])
    blocks = response['Blocks']

    key_map = {}
    value_map = {}
    block_map = {}
    for block in blocks:
        block_id = block['Id']
        block_map[block_id] = block
        if block['BlockType'] == "KEY_VALUE_SET":
            if 'KEY' in block['EntityTypes']:
                key_map[block_id] = block
            else:
                value_map[block_id] = block

    return key_map, value_map, block_map

def get_image_byte(filename):
    with open(filename, 'rb') as imagefile:
        return imagefile.read()

def save_file(extracted_text):
    file_type = [('Word Document', '*.docx')]
    filename = asksaveasfilename(
        filetypes=file_type, defaultextension=file_type)
    if filename:
        doc = docx.Document()
        doc.add_paragraph(extracted_text)
        doc.save(filename)

def ShowBoundingBox(draw,box,width,height,boxColor):          
    left = width * box['Left']
    top = height * box['Top'] 
    draw.rectangle([left,top, left + (width * box['Width']), top +(height * box['Height'])],outline=boxColor)


my_window.mainloop()
